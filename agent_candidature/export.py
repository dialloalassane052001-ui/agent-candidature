"""Export de la lettre de motivation en plusieurs formats : TXT, Word (.docx), PDF.

Chaque fonction prend le texte de la lettre et renvoie des `bytes`, prets a etre
proposes au telechargement (Streamlit) ou ecrits sur disque (CLI).
"""

import io

# Caracteres typographiques absents du latin-1 (police PDF standard) : on les
# remplace par un equivalent ASCII pour que le PDF ne plante pas sur les accents.
_REMPLACEMENTS = {
    "’": "'",  # apostrophe courbe
    "‘": "'",
    "“": '"',  # guillemets courbes
    "”": '"',
    "–": "-",  # tiret demi-cadratin
    "—": "-",  # tiret cadratin
    "…": "...",  # points de suspension
    " ": " ",  # espace insecable
    "•": "-",  # puce
}


def _assainir_latin1(texte: str) -> str:
    """Rend un texte compatible avec la police PDF de base (latin-1)."""
    for source, cible in _REMPLACEMENTS.items():
        texte = texte.replace(source, cible)
    # Filet de securite : tout caractere hors latin-1 est retire proprement.
    return texte.encode("latin-1", errors="ignore").decode("latin-1")


def lettre_vers_txt(texte: str) -> bytes:
    """Lettre au format texte brut (UTF-8)."""
    return texte.encode("utf-8")


def lettre_vers_docx(texte: str, titre: str = "Lettre de motivation") -> bytes:
    """Lettre au format Word (.docx), un paragraphe par bloc de texte."""
    from docx import Document
    from docx.shared import Pt

    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for bloc in texte.split("\n\n"):
        bloc = bloc.strip()
        if not bloc:
            continue
        paragraphe = document.add_paragraph()
        # Les retours a la ligne simples deviennent des sauts de ligne internes.
        lignes = bloc.split("\n")
        for i, ligne in enumerate(lignes):
            if i > 0:
                paragraphe.add_run().add_break()
            paragraphe.add_run(ligne)

    tampon = io.BytesIO()
    document.save(tampon)
    return tampon.getvalue()


def lettre_vers_pdf(texte: str, titre: str = "Lettre de motivation") -> bytes:
    """Lettre au format PDF (mise en page A4 simple et lisible)."""
    from fpdf import FPDF

    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_margins(left=25, top=20, right=25)
    pdf.set_font("Helvetica", size=11)

    for bloc in _assainir_latin1(texte).split("\n\n"):
        bloc = bloc.strip()
        if not bloc:
            continue
        # multi_cell gere le retour a la ligne automatique dans la largeur de page.
        pdf.multi_cell(0, 6, bloc)
        pdf.ln(3)

    sortie = pdf.output()
    # fpdf2 renvoie un bytearray : on normalise en bytes.
    return bytes(sortie)
